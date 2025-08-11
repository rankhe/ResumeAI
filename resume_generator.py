"""
简历生成模块
负责生成最终的简历文件（PDF/DOCX格式）
"""

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
import os
from typing import Dict

class ResumeGenerator:
    def __init__(self):
        pass
    
    def generate_resume(self, optimized_content: str, format_type: str = "pdf", output_path: str = "resume_output") -> str:
        """
        生成简历文件
        
        Args:
            optimized_content: 优化后的简历内容
            format_type: 文件格式 ('pdf' 或 'docx')
            output_path: 输出文件路径（不含扩展名）
            
        Returns:
            生成的文件路径
        """
        if format_type.lower() == "pdf":
            return self._generate_pdf(optimized_content, f"{output_path}.pdf")
        elif format_type.lower() == "docx":
            return self._generate_docx(optimized_content, f"{output_path}.docx")
        else:
            raise ValueError(f"不支持的文件格式: {format_type}")
    
    def _generate_pdf(self, content: str, output_path: str) -> str:
        """
        生成PDF格式简历
        
        Args:
            content: 简历内容
            output_path: 输出文件路径
            
        Returns:
            生成的文件路径
        """
        try:
            # 创建PDF文档
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            story = []
            
            # 获取样式表
            styles = getSampleStyleSheet()
            title_style = styles['Title']
            heading_style = styles['Heading2']
            normal_style = styles['Normal']
            
            # 解析内容并添加到PDF
            lines = content.split('\n')
            for line in lines:
                if line.startswith("===") or line.startswith("---"):
                    # 分隔线，跳过或添加空格
                    story.append(Spacer(1, 0.2*inch))
                elif "优化后的简历" in line:
                    story.append(Paragraph("优化后的简历", title_style))
                    story.append(Spacer(1, 0.2*inch))
                elif "针对职位:" in line:
                    story.append(Paragraph(line, heading_style))
                elif "公司:" in line:
                    story.append(Paragraph(line, heading_style))
                elif "优化说明:" in line:
                    story.append(Paragraph("优化说明", heading_style))
                elif line.strip() and not line.isspace():
                    story.append(Paragraph(line, normal_style))
                else:
                    story.append(Spacer(1, 0.1*inch))
            
            # 构建PDF
            doc.build(story)
            return output_path
        except Exception as e:
            raise Exception(f"PDF生成失败: {str(e)}")
    
    def _generate_docx(self, content: str, output_path: str) -> str:
        """
        生成DOCX格式简历
        
        Args:
            content: 简历内容
            output_path: 输出文件路径
            
        Returns:
            生成的文件路径
        """
        try:
            # 创建文档
            doc = Document()
            
            # 添加标题
            doc.add_heading('优化后的简历', 0)
            
            # 解析内容并添加到文档
            lines = content.split('\n')
            in_optimization_notes = False
            
            for line in lines:
                if line.startswith("===") or line.startswith("---"):
                    # 跳过分隔线
                    continue
                elif "针对职位:" in line:
                    doc.add_heading('职位信息', level=1)
                    doc.add_paragraph(line)
                elif "公司:" in line:
                    doc.add_paragraph(line)
                elif "优化说明:" in line:
                    doc.add_heading('优化说明', level=1)
                    in_optimization_notes = True
                elif line.strip() and not line.isspace():
                    if in_optimization_notes and line.startswith(("1.", "2.", "3.", "4.", "5.")):
                        # 处理列表项
                        doc.add_paragraph(line, style='List Number')
                    else:
                        doc.add_paragraph(line)
            
            # 保存文档
            doc.save(output_path)
            return output_path
        except Exception as e:
            raise Exception(f"DOCX生成失败: {str(e)}")
    
    def generate_ats_friendly_resume(self, resume_data: Dict, output_path: str = "ats_friendly_resume.pdf") -> str:
        """
        生成ATS友好的简历
        
        Args:
            resume_data: 简历数据
            output_path: 输出文件路径
            
        Returns:
            生成的文件路径
        """
        try:
            # 创建ATS友好的PDF
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            story = []
            
            # 使用简单样式确保ATS友好
            styles = getSampleStyleSheet()
            normal_style = styles['Normal']
            heading_style = ParagraphStyle(
                'Heading',
                parent=normal_style,
                fontSize=14,
                spaceAfter=12,
                bold=True
            )
            
            # 添加联系信息
            contact_info = resume_data.get("contact_info", {})
            if contact_info:
                name = contact_info.get("name", "")
                if name:
                    story.append(Paragraph(name, heading_style))
                
                email = contact_info.get("email", "")
                phone = contact_info.get("phone", "")
                contact_line = f"{email} | {phone}" if email and phone else email or phone
                if contact_line:
                    story.append(Paragraph(contact_line, normal_style))
                    story.append(Spacer(1, 0.2*inch))
            
            # 添加技能部分
            skills = resume_data.get("skills", [])
            if skills:
                story.append(Paragraph("技能", heading_style))
                skills_text = ", ".join(skills)
                story.append(Paragraph(skills_text, normal_style))
                story.append(Spacer(1, 0.2*inch))
            
            # 添加工作经验
            work_experience = resume_data.get("work_experience", [])
            if work_experience:
                story.append(Paragraph("工作经验", heading_style))
                for exp in work_experience:
                    # 简化处理，实际应用中需要更详细的信息
                    exp_text = f"工作经验 ({exp.get('start_year', '')} - {exp.get('end_year', '')})"
                    story.append(Paragraph(exp_text, normal_style))
                story.append(Spacer(1, 0.2*inch))
            
            # 添加教育背景
            education = resume_data.get("education", [])
            if education:
                story.append(Paragraph("教育背景", heading_style))
                for edu in education:
                    edu_text = f"{edu.get('institution', '')} - {edu.get('degree', '')}"
                    story.append(Paragraph(edu_text, normal_style))
            
            # 构建PDF
            doc.build(story)
            return output_path
        except Exception as e:
            raise Exception(f"ATS友好简历生成失败: {str(e)}")

# 使用示例
if __name__ == "__main__":
    generator = ResumeGenerator()
    # 示例用法
    # file_path = generator.generate_resume("优化后的简历内容", "pdf")
    # print(f"简历已生成: {file_path}")