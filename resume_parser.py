"""
简历解析模块
负责解析用户上传的简历文件，提取关键信息
"""

import PyPDF2
import pdfplumber
from docx import Document
import re
from typing import Dict, List
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize

# 下载必要的NLTK数据（首次运行时需要）
try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('corpora/stopwords')
except LookupError:
    print("首次运行需要下载NLTK数据，请稍候...")
    nltk.download('punkt')
    nltk.download('stopwords')

class ResumeParser:
    def __init__(self):
        # 中文停用词
        self.chinese_stopwords = {
            '的', '了', '在', '是', '我', '有', '和', '就', '不', '人', '都', '一', '一个', '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看', '好', '自己', '这'
        }
        
        # 英文停用词
        self.english_stopwords = set(stopwords.words('english'))
        
        # 工作经验相关关键词
        self.exp_keywords = [
            '工作经历', '工作经验', '工作背景', '职业经历', 'employment', 'experience', 'work history', 'professional experience', 'career history'
        ]
        
        # 教育背景相关关键词
        self.edu_keywords = [
            '教育背景', '教育经历', '学历', '学位', '毕业院校', 'education', 'degree', 'academic background', 'university', 'college'
        ]
        
        # 技能相关关键词
        self.skill_keywords = [
            '技能', '技术', '能力', '专长', 'skills', 'technologies', 'tools', 'abilities', 'expertise'
        ]
    
    def parse_resume(self, file_path: str, file_type: str) -> Dict:
        """
        解析简历文件
        
        Args:
            file_path: 简历文件路径
            file_type: 文件类型 ('pdf'、'docx' 或 'txt')
            
        Returns:
            包含简历信息的字典
        """
        if file_type.lower() == 'pdf':
            return self._parse_pdf(file_path)
        elif file_type.lower() == 'docx':
            return self._parse_docx(file_path)
        elif file_type.lower() == 'txt':
            return self._parse_txt(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {file_type}")
    
    def _parse_pdf(self, file_path: str) -> Dict:
        """
        解析PDF格式简历
        """
        resume_data = {
            "text": "",
            "contact_info": {},
            "work_experience": [],
            "education": [],
            "skills": [],
            "projects": [],
            "certifications": []
        }
        
        try:
            # 使用pdfplumber提取文本
            with pdfplumber.open(file_path) as pdf:
                text = ""
                for page in pdf.pages:
                    text += page.extract_text() or ""
                resume_data["text"] = text
            
            # 提取关键信息
            resume_data["contact_info"] = self._extract_contact_info(resume_data["text"])
            resume_data["work_experience"] = self._extract_work_experience(resume_data["text"])
            resume_data["education"] = self._extract_education(resume_data["text"])
            resume_data["skills"] = self._extract_skills(resume_data["text"])
            resume_data["projects"] = self._extract_projects(resume_data["text"])
            resume_data["certifications"] = self._extract_certifications(resume_data["text"])
            
            return resume_data
        except Exception as e:
            raise Exception(f"PDF简历解析失败: {str(e)}")
    
    def _parse_docx(self, file_path: str) -> Dict:
        """
        解析DOCX格式简历
        """
        try:
            doc = Document(file_path)
            # 提取所有段落和表格中的文本
            paragraphs_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # 提取表格中的文本
            tables_text = ""
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        tables_text += cell.text + "\n"
            
            text = paragraphs_text + "\n" + tables_text
            
            resume_data = {
                "text": text,
                "contact_info": self._extract_contact_info(text),
                "work_experience": self._extract_work_experience(text),
                "education": self._extract_education(text),
                "skills": self._extract_skills(text),
                "projects": self._extract_projects(text),
                "certifications": self._extract_certifications(text)
            }
            
            return resume_data
        except Exception as e:
            raise Exception(f"DOCX简历解析失败: {str(e)}")
    
    def _parse_txt(self, file_path: str) -> Dict:
        """
        解析TXT格式简历
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            resume_data = {
                "text": text,
                "contact_info": self._extract_contact_info(text),
                "work_experience": self._extract_work_experience(text),
                "education": self._extract_education(text),
                "skills": self._extract_skills(text),
                "projects": self._extract_projects(text),
                "certifications": self._extract_certifications(text)
            }
            
            return resume_data
        except UnicodeDecodeError:
            # 如果UTF-8解码失败，尝试其他编码
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    text = f.read()
                
                resume_data = {
                    "text": text,
                    "contact_info": self._extract_contact_info(text),
                    "work_experience": self._extract_work_experience(text),
                    "education": self._extract_education(text),
                    "skills": self._extract_skills(text),
                    "projects": self._extract_projects(text),
                    "certifications": self._extract_certifications(text)
                }
                
                return resume_data
            except Exception as e:
                raise Exception(f"TXT简历解析失败: {str(e)}")
        except Exception as e:
            raise Exception(f"TXT简历解析失败: {str(e)}")
    
    def _extract_contact_info(self, text: str) -> Dict:
        """
        从简历文本中提取联系信息
        """
        contact_info = {}
        
        # 提取邮箱 (更精确的正则表达式)
        email_pattern = r'\b[A-Za-z0-9](?:[A-Za-z0-9._-]*[A-Za-z0-9])?@[A-Za-z0-9](?:[A-Za-z0-9.-]*[A-Za-z0-9])?\.[A-Za-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            contact_info["email"] = emails[0]
        
        # 提取电话号码 (支持多种格式)
        phone_patterns = [
            r'(?:\+?86[-\s]?)?(?:1[3-9]\d{9})',  # 中国手机号
            r'(?:\+?86[-\s]?)?(?:\d{3,4}[-\s]?\d{7,8})',  # 中国固话
            r'(?:\+?1[-\s]?)?\(?[2-9]\d{2}\)?[-\s]?[2-9]\d{2}[-\s]?\d{4}',  # 美国号码
            r'\+?\d{1,4}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}'  # 通用格式
        ]
        
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                contact_info["phone"] = phones[0].strip()
                break
        
        # 提取姓名（使用更智能的方法）
        # 假设姓名在简历开头附近，且为2-4个中文字符或常见的英文姓名格式
        lines = text.split('\n')
        for i, line in enumerate(lines[:10]):  # 检查前10行
            line = line.strip()
            # 中文姓名（2-4个汉字）
            chinese_name_pattern = r'^[\u4e00-\u9fa5]{2,4}$'
            # 英文姓名（名+姓，可能有中间名）
            english_name_pattern = r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]*\.?)?(?:\s+[A-Z][a-z]+){1,2}$'
            
            if re.match(chinese_name_pattern, line) or re.match(english_name_pattern, line):
                contact_info["name"] = line
                break
        
        # 如果还没找到姓名，使用简化的方法
        if "name" not in contact_info:
            name_pattern = r'^([A-Z][a-z]+\s+){1,3}[A-Z][a-z]+'
            names = re.findall(name_pattern, text)
            if names:
                contact_info["name"] = names[0].strip()
            else:
                # 尝试查找单独的一行，可能包含姓名
                lines = text.split('\n')
                for line in lines[:5]:
                    line = line.strip()
                    if 2 <= len(line) <= 30 and not re.search(r'[@:0-9]', line):  # 不包含@、:、数字
                        if re.search(r'[\u4e00-\u9fa5]|[A-Z][a-z]', line):  # 包含中文或英文
                            contact_info["name"] = line
                            break
        
        # 提取LinkedIn或其他社交媒体
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        linkedin_matches = re.findall(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_matches:
            contact_info["linkedin"] = "https://www." + linkedin_matches[0]
        
        github_pattern = r'github\.com/[A-Za-z0-9_-]+'
        github_matches = re.findall(github_pattern, text, re.IGNORECASE)
        if github_matches:
            contact_info["github"] = "https://" + github_matches[0]
        
        return contact_info
    
    def _extract_work_experience(self, text: str) -> List[Dict]:
        """
        提取工作经验
        """
        experiences = []
        
        # 找到工作经验部分
        exp_section_start = -1
        exp_section_text = ""
        
        for keyword in self.exp_keywords:
            if keyword.lower() in text.lower():
                exp_section_start = text.lower().find(keyword.lower())
                break
        
        if exp_section_start != -1:
            # 找到工作经验部分的结束位置
            exp_section_end = len(text)
            # 查找下一个主要部分的开始位置
            next_section_keywords = self.edu_keywords + self.skill_keywords
            for keyword in next_section_keywords:
                keyword_pos = text.lower().find(keyword.lower(), exp_section_start + 1)
                if keyword_pos != -1 and keyword_pos < exp_section_end:
                    exp_section_end = keyword_pos
            
            exp_section_text = text[exp_section_start:exp_section_end]
        else:
            # 如果没有找到明确的工作经验部分，使用整个文本
            exp_section_text = text
        
        # 从工作经验部分提取具体的工作经历
        # 查找公司名称和职位的模式
        # 这里使用简化的模式，实际应用中可能需要更复杂的NLP技术
        
        # 查找时间范围
        date_patterns = [
            r'(\d{4}(?:\s*[-—–~]\s*\d{4}|(?:\s*[-—–~]\s*)?至今|present|current))',
            r'(\d{4}\s*年\s*\d{1,2}\s*月\s*[-—–~]\s*\d{4}\s*年\s*\d{1,2}\s*月)',
        ]
        
        for pattern in date_patterns:
            matches = re.finditer(pattern, exp_section_text)
            for match in matches:
                # 对于每个时间范围，查找前面的公司和职位信息
                start_pos = max(0, match.start() - 200)  # 向前查找200个字符
                context = exp_section_text[start_pos:match.end()]
                
                # 简化处理：提取一行包含时间的内容
                lines = context.split('\n')
                for line in lines:
                    if match.group(1) in line and len(line.strip()) > 5:
                        # 提取公司名称（启发式方法）
                        company_pattern = r'^[^\d\n]{3,30}'
                        company_match = re.search(company_pattern, line)
                        company = company_match.group(0).strip() if company_match else "未知公司"
                        
                        experiences.append({
                            "company": company,
                            "title": "未知职位",
                            "duration": match.group(1),
                            "description": line.strip()
                        })
                        break
        
        # 如果没有通过时间范围找到，尝试其他方法
        if not experiences:
            # 查找包含公司特征的行（以有限公司、公司等结尾）
            company_lines = re.findall(r'.*(?:有限公司|公司|Company|Corp|LLC).*', exp_section_text)
            for line in company_lines[:5]:  # 限制数量
                experiences.append({
                    "company": line.strip(),
                    "title": "未知职位",
                    "duration": "未知时间",
                    "description": line.strip()
                })
        
        return experiences
    
    def _extract_education(self, text: str) -> List[Dict]:
        """
        提取教育背景
        """
        education = []
        
        # 找到教育背景部分
        edu_section_start = -1
        edu_section_text = ""
        
        for keyword in self.edu_keywords:
            if keyword.lower() in text.lower():
                edu_section_start = text.lower().find(keyword.lower())
                break
        
        if edu_section_start != -1:
            # 找到教育背景部分的结束位置
            edu_section_end = len(text)
            # 查找下一个主要部分的开始位置
            next_section_keywords = self.skill_keywords + ['工作经历', '项目经验']
            for keyword in next_section_keywords:
                keyword_pos = text.lower().find(keyword.lower(), edu_section_start + 1)
                if keyword_pos != -1 and keyword_pos < edu_section_end:
                    edu_section_end = keyword_pos
            
            edu_section_text = text[edu_section_start:edu_section_end]
        else:
            # 如果没有找到明确的教育背景部分，使用整个文本
            edu_section_text = text
        
        # 查找学校名称
        # 包含大学、学院等关键词的行
        university_patterns = [
            r'.*(大学|学院|university|college|school).*',
            r'.*(清华|北大|复旦|交大|浙大).*',  # 一些知名大学简称
        ]
        
        found_education = False
        for pattern in university_patterns:
            universities = re.findall(pattern, edu_section_text, re.IGNORECASE)
            for uni_match in universities:
                if isinstance(uni_match, tuple):
                    uni_text = ''.join(uni_match)
                else:
                    uni_text = uni_match
                
                # 获取匹配行的完整文本
                lines = edu_section_text.split('\n')
                for line in lines:
                    if uni_text in line and len(line.strip()) > 4:
                        # 尝试提取学位信息
                        degree_patterns = ['学士', '硕士', '博士', '本科', '研究生', 'Bachelor', 'Master', 'PhD', 'Degree']
                        degree = "未知学位"
                        for deg in degree_patterns:
                            if deg in line:
                                degree = deg
                                break
                        
                        education.append({
                            "institution": line.strip(),
                            "degree": degree,
                            "major": "未知专业",
                            "graduation_year": "未知"
                        })
                        found_education = True
                        break
                if found_education:
                    break
            if found_education:
                break
        
        return education
    
    def _extract_skills(self, text: str) -> List[str]:
        """
        提取技能列表
        """
        # 常见技能关键词（扩展列表）
        common_skills = [
            'Python', 'Java', 'JavaScript', 'C++', 'C#', 'SQL', 'PHP', 'Ruby', 'Go', 'Rust',
            'React', 'Vue', 'Angular', 'Node.js', 'Express', 'Django', 'Flask', 'Spring', 'ASP.NET',
            'Docker', 'Kubernetes', 'AWS', 'Azure', 'GCP', 'Linux', 'Windows', 'macOS',
            '数据分析', '机器学习', '深度学习', '人工智能', 'TensorFlow', 'PyTorch', 'Keras',
            'Git', 'Jenkins', 'CI/CD', 'DevOps', 'Agile', 'Scrum', 'JIRA',
            'MySQL', 'PostgreSQL', 'MongoDB', 'Redis', 'Oracle', 'SQL Server',
            'HTML', 'CSS', 'Bootstrap', 'jQuery', 'TypeScript', 'RESTful', 'API',
            '数据分析', '数据科学', '统计学', 'R语言', 'Tableau', 'Power BI'
        ]
        
        found_skills = []
        text_lower = text.lower()
        
        for skill in common_skills:
            if skill.lower() in text_lower:
                found_skills.append(skill)
        
        # 查找技能部分并提取更多技能
        skill_section_start = -1
        for keyword in self.skill_keywords:
            if keyword.lower() in text_lower:
                skill_section_start = text_lower.find(keyword.lower())
                break
        
        if skill_section_start != -1:
            # 定义技能部分的结束位置
            skill_section_end = len(text)
            next_sections = ['工作经历', '工作经验', '教育背景', '项目经验', '证书']
            for section in next_sections:
                pos = text_lower.find(section.lower(), skill_section_start + 1)
                if pos != -1 and pos < skill_section_end:
                    skill_section_end = pos
            
            skill_section_text = text[skill_section_start:skill_section_end]
            
            # 提取列表项和技能关键词
            list_items = re.findall(r'[•\-\*\+]\s*([^\n]+)', skill_section_text)
            for item in list_items:
                item_clean = item.strip()
                # 移除停用词，只保留有意义的技能词
                if len(item_clean) > 1 and item_clean not in self.chinese_stopwords and item_clean.lower() not in self.english_stopwords:
                    # 检查是否已经添加
                    if item_clean not in found_skills:
                        found_skills.append(item_clean)
            
            # 提取用逗号、分号或斜杠分隔的技能
            # 查找技能部分中的技能列表
            skill_lines = skill_section_text.split('\n')
            for line in skill_lines:
                if re.search(r'[;、,/]', line):  # 包含分隔符
                    skills_in_line = re.split(r'[;、,/]', line)
                    for skill in skills_in_line:
                        skill_clean = skill.strip()
                        if len(skill_clean) > 1 and skill_clean not in found_skills:
                            found_skills.append(skill_clean)
        
        # 去重并过滤空值
        found_skills = list(filter(None, list(set(found_skills))))
        
        return found_skills
    
    def _extract_projects(self, text: str) -> List[Dict]:
        """
        提取项目经验
        """
        projects = []
        current_project = None  # 在函数开始时初始化变量
        
        # 查找项目相关关键词
        project_keywords = ['项目经验', '项目经历', '项目背景', 'projects', 'project experience']
        project_section_start = -1
        
        text_lower = text.lower()
        for keyword in project_keywords:
            if keyword.lower() in text_lower:
                project_section_start = text_lower.find(keyword.lower())
                break
        
        if project_section_start != -1:
            # 定义项目部分的结束位置
            project_section_end = len(text)
            next_sections = ['工作经历', '教育背景', '技能', '证书']
            for section in next_sections:
                pos = text_lower.find(section.lower(), project_section_start + 1)
                if pos != -1 and pos < project_section_end:
                    project_section_end = pos
            
            project_section_text = text[project_section_start:project_section_end]
            
            # 提取项目名称和描述
            # 查找项目标题（通常是加粗或以"项目"开头的行）
            project_lines = project_section_text.split('\n')
            
            for line in project_lines:
                line = line.strip()
                if not line:
                    continue
                
                # 如果行以"项目"开头，或者看起来像项目名称
                if re.match(r'(项目.*[:：]|.*[Pp]roject.*)', line):
                    if current_project:
                        projects.append(current_project)
                    
                    current_project = {
                        "name": line,
                        "description": "",
                        "technologies": [],
                        "duration": "未知"
                    }
                elif current_project:
                    # 添加到当前项目描述中
                    if len(line) > 10:  # 过滤太短的行
                        current_project["description"] += line + " "
        
        # 如果有未添加的最后一个项目
        if current_project:
            projects.append(current_project)
        
        return projects
    
    def _extract_certifications(self, text: str) -> List[Dict]:
        """
        提取证书信息
        """
        certifications = []
        
        # 查找证书相关关键词
        cert_keywords = ['证书', '认证', 'certificates', 'certifications', 'credentials']
        cert_section_start = -1
        
        text_lower = text.lower()
        for keyword in cert_keywords:
            if keyword.lower() in text_lower:
                cert_section_start = text_lower.find(keyword.lower())
                break
        
        if cert_section_start != -1:
            # 定义证书部分的结束位置
            cert_section_end = len(text)
            next_sections = ['工作经历', '教育背景', '技能', '项目经验']
            for section in next_sections:
                pos = text_lower.find(section.lower(), cert_section_start + 1)
                if pos != -1 and pos < cert_section_end:
                    cert_section_end = pos
            
            cert_section_text = text[cert_section_start:cert_section_end]
            
            # 查找证书名称
            # 启发式方法：查找看起来像证书名称的行
            cert_lines = cert_section_text.split('\n')
            for line in cert_lines:
                line = line.strip()
                # 证书名称通常包含认证机构名称（如Cisco、Microsoft、Oracle等）
                cert_keywords = ['认证', '证书', 'Certified', 'Certificate']
                if any(keyword in line for keyword in cert_keywords) or \
                   any(org in line for org in ['Cisco', 'Microsoft', 'Oracle', 'AWS', 'Google', 'PMP', 'CCNA', 'CCNP']):
                    if len(line) > 5:  # 过滤太短的行
                        certifications.append({
                            "name": line,
                            "authority": "未知认证机构",
                            "date": "未知时间"
                        })
        
        return certifications

# 使用示例
if __name__ == "__main__":
    parser = ResumeParser()
    # 示例用法
    # resume_data = parser.parse_resume("resume.pdf", "pdf")
    # print(resume_data)